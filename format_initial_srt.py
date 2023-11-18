import PySimpleGUI as sg
import srt
import os
import sys
#import openpyxl
import csv
#from openpyxl import load_workbook
#from openpyxl.worksheet.table import Table
from docx import Document
from docx.shared import Cm
from bisect import bisect_left
from strictyaml import load, Map, Str, Int, as_document


def popup_error(message, title=None):
    layout = [
        [sg.Text(message,background_color=main_background_color,font=main_font)],
        [sg.Button('Retry',button=second_background_color,font=main_font)]
    ]

    window = sg.Window(title if title else message, layout)
    event, values = window.read(close=True)
    return

def create_formatted_srt():
    srt_file = load_srt()
    subs = list(srt.parse(srt_file))

    for idx, sub in enumerate(subs):
        if sub.content.count('\n') == 0:
            sub.content = sub.content + '\n'

    target_file_path = layout[2][0].get()

    target_file = open(target_file_path,'w', encoding='utf-8')
    target_file.write(srt.compose(subs))

    srt_file.close()
    target_file.close()

    sg.popup_ok("Formatted SRT successfully saved",title="Success",background_color=main_background_color, font=main_font)

def create_script_for_voice_over():
    srt_file = load_srt()
    subs = list(srt.parse(srt_file))

    target_file_path = layout[3][0].get()

    script = Document()

    previous_speaker = ''

    speaker_timecodes = load_speaker_timecode_csv()
    speaker_start_timecodes_as_float = [float(key[0]) for key in speaker_timecodes]
    speaker_end_timecodes_as_float = [float(key[1]) for key in speaker_timecodes]

    for sub in subs:
        closest_speaker_index = closest(speaker_start_timecodes_as_float, speaker_end_timecodes_as_float, sub)
        speaker = speaker_timecodes[closest_speaker_index][2]

        if previous_speaker != speaker:
            script.add_heading(speaker, level=1)
            previous_speaker = speaker

        script.add_paragraph(sub.content.replace("\n"," "))

    try:
        script.save(target_file_path)

        sg.popup_ok("Script DOCX successfully saved",title="Success",background_color=main_background_color, font=main_font)
        return
    except PermissionError:
        popup_error("Could not save file. Is the target file still open in Word?",title="Error")

def closest(start_timecodes, end_timecodes, target_sub):
    target_start = target_sub.start.total_seconds()
    target_end = target_sub.end.total_seconds()
    timecode_length = len(start_timecodes)

    # Find best matching timecode for start of subtitle
    pos_start = bisect_left(start_timecodes, target_start)

    if pos_start == 0:
        return pos_start
    if pos_start == timecode_length:
        return -1

    # Check speaker timecodes before, on and after the best matching timecode
    before_index = min(max(0, pos_start - 1), timecode_length - 1)
    before_start = start_timecodes[before_index]
    before_end = end_timecodes[before_index]
    before_from_timecode = max(before_start, target_start)
    before_to_timecode = min(before_end, target_end)
    before_coverage = before_to_timecode - before_from_timecode

    current_index = min(max(0, pos_start), timecode_length - 1)
    current_start = start_timecodes[current_index]
    current_end = end_timecodes[current_index]
    current_from_timecode = max(current_start, target_start)
    current_to_timecode = min(current_end, target_end)
    current_coverage = current_to_timecode - current_from_timecode

    after_index = min(max(0, pos_start + 1), timecode_length - 1)
    after_start = start_timecodes[after_index]
    after_end = end_timecodes[after_index]
    after_from_timecode = max(after_start, target_start)
    after_to_timecode = min(after_end, target_end)
    after_coverage = after_to_timecode - after_from_timecode

    #print(f"---")
    #print(f"Target start/end: {target_start}, {target_end}")
    #print(f"Before from/to/coverage: {before_from_timecode}, {before_to_timecode}, {before_coverage}")
    #print(f"Current from/to/coverage: {current_from_timecode}, {current_to_timecode}, {current_coverage}")
    #print(f"After from/to/coverage: {after_from_timecode}, {after_to_timecode}, {after_coverage}")

    # Get timecodes for largest coverage
    coverages = [before_coverage, current_coverage, after_coverage]
    coverages.sort()
    largest_coverage = coverages[-1]

    if before_coverage == largest_coverage:
        #print(f"Returning before_index: {before_index}")
        return before_index
    elif current_coverage == largest_coverage:
        #print(f"Returning current_index: {current_index}")
        return current_index
    else:
        #print(f"Returning after_index: {after_index}")
        return after_index

def load_srt():
    try:
        file_path = layout[0][1].get()
        srt_file = open(file_path, 'r', encoding='utf-8')

        return srt_file
    except FileNotFoundError:
        popup_error("SRT input file was not found. Is the path correct?",title="Error")

def load_speaker_timecode_csv():
    try:
        file_path = layout[1][1].get()
        speaker_timecode_csv_file = open(file_path, 'r', encoding='utf-8')

        with open(file_path, mode='r') as csv_file:
            reader = csv.reader(csv_file, delimiter=';', quotechar='"')

            speaker_timecodes = [[rows[3], rows[4], rows[2]] for rows in reader]

            # Delete header line
            speaker_timecodes = speaker_timecodes[1:]

        return speaker_timecodes
    except FileNotFoundError:
        popup_error("Speaker Timecode CSV input file was not found. Is the path correct?",title="Error")
    except IndexError:
        popup_error("Speaker Timecode CSV could not be read. Make sure to use a semicolon (;) as a separator and quotation marks for text.")

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

def load_config():
    config_file_path = "config/subtitle_tool_config.yaml"
    schema = Map({
        "main_background_color": Str(),
        "second_background_color": Str(),
        "font_family": Str(),
        "font_size": Int()})
    
    try:
        with open(config_file_path, mode='r') as config_file:
            config_data = config_file.read()
            config = load(config_data, schema).data
            config_file.close()
    except FileNotFoundError:
        config = create_default_config(schema, config_file_path)
        
    return config

def create_default_config(schema: Map, config_file_path):
    yaml = as_document(schema = schema, data={
        "main_background_color": "#434243",
        "second_background_color": "#faa61a",
        "font_family": "Lato",
        "font_size": 12
    })
    with open(config_file_path, mode='w') as config_file:
        config_file.write(yaml.as_yaml())
        config_file.close()
    
    return yaml.data

config = load_config()

main_background_color = config['main_background_color']
second_background_color = config['second_background_color']
main_font_family = config['font_family']
main_font = (main_font_family,config['font_size'])

layout = [
    [
        sg.Text("Input Subtitle SRT file:",background_color=main_background_color, font=main_font),
        sg.In(font=main_font), sg.FileBrowse(key="-BROWSE_INPUT", button_color=second_background_color, font=main_font)
    ],
    [
        sg.Text("Input Speaker Timecodes CSV file:",background_color=main_background_color, font=main_font),
        sg.In(font=main_font), sg.FileBrowse(key="-BROWSE_INPUT", button_color=second_background_color, font=main_font)
    ],
    [
        sg.In(key="-SRT_TARGET_FILLED",visible=False,enable_events=True),
        sg.FileSaveAs("Save SRT file as formatted SRT file",default_extension=".srt",file_types=(("SRT files",".srt"),),button_color=second_background_color, font=main_font),
    ],
    [
        sg.In(key="-SCRIPT_TARGET_FILLED",visible=False,enable_events=True),
        sg.FileSaveAs("Merge SRT file and speaker timecodes into a script",default_extension=".docx",file_types=(("Word files (docx)",".docx"),),button_color=second_background_color, font=main_font),
    ],
#    [
#        sg.In(key="-XLSX_ADD_SRT",visible=False,enable_events=True),
#        sg.FileSaveAs("Add SRT to XLSX",default_extension=".xlsx",file_types=(("Excel files (xlsx)",".xlsx"),),button_color=second_background_color, font=main_font),
#    ]
]

window = sg.Window("Subtitle Tools for The Week",layout,icon=resource_path("theweek.ico"), background_color=main_background_color) 

while True:
    event, values = window.read()

    if event == "-SRT_TARGET_FILLED":
        create_formatted_srt()

    if event == "-SCRIPT_TARGET_FILLED":
        create_script_for_voice_over()

#    if event == "-XLSX_ADD_SRT":
#        add_srt_to_xlsx()
        
    # End program if user closes window
    if event == sg.WIN_CLOSED:
        break

window.close()

